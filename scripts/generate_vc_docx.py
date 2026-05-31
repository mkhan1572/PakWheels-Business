from docx import Document
import subprocess
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
md_path = ROOT / "Documentation" / "Version_Control_Document.md"
docx_path = ROOT / "Documentation" / "Version_Control_Document.docx"

doc = Document()
doc.add_heading('Version Control Document', level=1)

if md_path.exists():
    with md_path.open('r', encoding='utf-8') as f:
        doc.add_paragraph(f.read())

# Append git log
try:
    git_log = subprocess.check_output(["git", "log", "--pretty=format:%h %ad %s", "--date=short"], cwd=ROOT, text=True)
    doc.add_heading('Git Commit History', level=2)
    for line in git_log.splitlines():
        doc.add_paragraph(line)
except Exception as e:
    doc.add_paragraph(f"Could not read git log: {e}")

doc.save(docx_path)
print(f"Wrote {docx_path}")
